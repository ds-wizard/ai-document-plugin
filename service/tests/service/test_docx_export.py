import io
from dataclasses import dataclass

import docx
import pytest
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from ai_document_plugin_service.service.docx_export import markdown_to_docx


@dataclass(frozen=True)
class _NumberedParagraph:
    text: str
    num_id: int
    level: int


def _render(markdown: str) -> Document:
    return docx.Document(io.BytesIO(markdown_to_docx(markdown)))


def _style_name(paragraph: Paragraph) -> str | None:
    return paragraph.style.name if paragraph.style is not None else None


def _numbering_of(paragraph: Paragraph) -> tuple[int, int] | None:
    """Return (numId, ilvl) for a numbered paragraph, or None when it is not numbered."""
    properties = paragraph._p.pPr  # noqa: SLF001
    if properties is None or properties.numPr is None:
        return None
    return int(properties.numPr.numId.val), int(properties.numPr.ilvl.val)


def _numbered_paragraphs(document: Document) -> list[_NumberedParagraph]:
    """Every numbered paragraph in document order, with the list it belongs to."""
    numbered = []
    for paragraph in document.paragraphs:
        numbering = _numbering_of(paragraph)
        if numbering is not None:
            numbered.append(_NumberedParagraph(paragraph.text, numbering[0], numbering[1]))
    return numbered


def test_headings_use_word_heading_styles() -> None:
    document = _render('# Title\n\n## Section\n\n###### Deep')

    styles = [_style_name(paragraph) for paragraph in document.paragraphs if paragraph.text]
    assert styles == ['Heading 1', 'Heading 2', 'Heading 6']
    assert [paragraph.text for paragraph in document.paragraphs if paragraph.text] == [
        'Title',
        'Section',
        'Deep',
    ]


def test_heading_deeper_than_six_is_clamped() -> None:
    # CommonMark caps ATX headings at six, but a clamp keeps a bad style name from being used.
    document = _render('###### Six')

    assert _style_name(document.paragraphs[0]) == 'Heading 6'


def test_inline_formatting_maps_to_runs() -> None:
    document = _render('Plain **bold** and *italic* and ~~struck~~ and `code`.')

    runs = {run.text: run for run in document.paragraphs[0].runs}
    assert runs['bold'].bold is True
    assert runs['italic'].italic is True
    assert runs['struck'].font.strike is True
    assert runs['code'].font.name == 'Consolas'
    # Formatting is only ever switched on, never explicitly off, so styles keep inheriting.
    assert runs['Plain '].bold is None


def test_heading_runs_do_not_disable_inherited_bold() -> None:
    document = _render('# Title with **bold**')

    assert all(run.bold in (None, True) for run in document.paragraphs[0].runs)


def test_soft_break_becomes_a_space_in_one_paragraph() -> None:
    document = _render('first\nsecond')

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert paragraphs == ['first second']


def test_link_becomes_a_hyperlink_relationship() -> None:
    document = _render('See [the guide](https://example.org/guide).')

    paragraph = document.paragraphs[0]
    hyperlinks = paragraph._p.findall(qn('w:hyperlink'))  # noqa: SLF001
    assert len(hyperlinks) == 1

    relationship_id = hyperlinks[0].get(qn('r:id'))
    assert paragraph.part.rels[relationship_id].target_ref == 'https://example.org/guide'
    assert paragraph.text.startswith('See ')
    assert 'the guide' in paragraph.text


def test_separate_ordered_lists_restart_numbering() -> None:
    document = _render('1. one\n2. two\n\nBetween.\n\n1. alpha\n2. beta')

    numbered = _numbered_paragraphs(document)
    assert [item.text for item in numbered] == ['one', 'two', 'alpha', 'beta']

    assert numbered[1].num_id == numbered[0].num_id
    # A shared numId is what would make the second list continue 3., 4. instead of restarting.
    assert numbered[2].num_id != numbered[0].num_id


def test_numbering_part_keeps_schema_element_order() -> None:
    # w:numbering requires every w:abstractNum to precede every w:num. Word refuses to open the
    # file if that is violated, and nothing else in the test suite would notice.
    document = _render('1. one\n\ntext\n\n- bullet\n\ntext\n\n1. two')

    tags = [child.tag.split('}')[1] for child in document.part.numbering_part.element]
    assert tags.count('abstractNum') == tags.count('num') >= 3
    assert max(i for i, tag in enumerate(tags) if tag == 'abstractNum') < min(
        i for i, tag in enumerate(tags) if tag == 'num'
    )


def test_allocated_numbering_ids_are_unique() -> None:
    document = _render('1. one\n\ntext\n\n1. two\n\ntext\n\n- three')

    num_ids = [element.get(qn('w:numId')) for element in document.part.numbering_part.element.findall(qn('w:num'))]
    assert len(num_ids) == len(set(num_ids))


def test_ordered_list_honours_explicit_start() -> None:
    document = _render('5. five\n6. six')

    numbering = document.part.numbering_part.element
    starts = [element.get(qn('w:val')) for element in numbering.iter(qn('w:start'))]
    assert '5' in starts


def test_nested_lists_use_increasing_levels() -> None:
    document = _render('- top\n    - nested\n        - deeper')

    assert [item.level for item in _numbered_paragraphs(document)] == [0, 1, 2]


def test_list_levels_beyond_the_maximum_are_capped() -> None:
    markdown = '\n'.join(f'{" " * (4 * depth)}- level {depth}' for depth in range(9))
    document = _render(markdown)

    assert max(item.level for item in _numbered_paragraphs(document)) == 5


def test_loose_list_item_numbers_only_its_first_paragraph() -> None:
    document = _render('1. first para\n\n   second para\n\n2. next item')

    assert [item.text for item in _numbered_paragraphs(document)] == ['first para', 'next item']


def test_table_renders_with_bold_header() -> None:
    document = _render('| Name | Size |\n| --- | ---: |\n| disk | 10 |\n| tape | 20 |')

    table = document.tables[0]
    assert len(table.rows) == 3
    assert len(table.columns) == 2
    assert table.cell(0, 0).text == 'Name'
    assert table.cell(2, 1).text == '20'
    assert all(run.bold for run in table.cell(0, 0).paragraphs[0].runs)
    assert not any(run.bold for run in table.cell(1, 0).paragraphs[0].runs)


def test_table_column_alignment_is_applied() -> None:
    document = _render('| L | R |\n| :--- | ---: |\n| a | b |')

    right_aligned = document.tables[0].cell(1, 1).paragraphs[0]
    assert right_aligned.alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_fenced_code_keeps_one_paragraph_per_line() -> None:
    document = _render('```python\nfirst = 1\nsecond = 2\n```')

    code_paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.runs and paragraph.runs[0].font.name == 'Consolas'
    ]
    assert code_paragraphs == ['first = 1', 'second = 2']


def test_blockquote_uses_the_quote_style() -> None:
    document = _render('> quoted text')

    assert any(
        _style_name(paragraph) == 'Quote' and paragraph.text == 'quoted text' for paragraph in document.paragraphs
    )


def test_horizontal_rule_becomes_a_bottom_border() -> None:
    document = _render('above\n\n---\n\nbelow')

    borders = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.pPr is not None and paragraph._p.pPr.find(qn('w:pBdr')) is not None  # noqa: SLF001
    ]
    assert len(borders) == 1


def test_raw_html_is_kept_as_literal_text_not_interpreted() -> None:
    document = _render('<script>alert(1)</script>')

    text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
    assert 'alert(1)' in text


def test_image_falls_back_to_alt_text() -> None:
    document = _render('![a diagram](https://example.org/x.png)')

    assert document.paragraphs[0].text == '[a diagram]'


def test_title_is_stored_as_document_metadata() -> None:
    content = markdown_to_docx('# Body', title='My Plan')
    document = docx.Document(io.BytesIO(content))

    assert document.core_properties.title == 'My Plan'


@pytest.mark.parametrize('markdown', ['', '   \n\n  '])
def test_empty_markdown_still_produces_a_readable_file(markdown: str) -> None:
    document = docx.Document(io.BytesIO(markdown_to_docx(markdown)))

    assert document.paragraphs == [] or all(not p.text for p in document.paragraphs)

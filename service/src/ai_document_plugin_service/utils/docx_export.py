"""Convert generated markdown into a Word (.docx) document.

The markdown produced by the pipeline is CommonMark plus the GitHub extensions the plugin
renders in the browser (tables, strikethrough, and bare URLs as links), so the parser here is
configured to match `remark-gfm` on the frontend. Linkify matters in practice: the pipeline
writes repository and dataset URLs as bare text rather than as `[label](url)`.

Everything maps onto Word's built-in styles, which keeps this module small, lets headings feed
Word's navigation pane and generated tables of contents, and lets an institution restyle the
output by editing the styles rather than this code. The trade is that Word's `List Number`
style carries a single counter, so consecutive numbered lists continue rather than restart --
see `test_docx_export.py`.

Hyperlinks are the one place that drops to raw OOXML, because python-docx exposes no API for
writing them.

Images and thematic breaks are not handled: the pipeline does not produce them. Neither breaks
the export if one does appear -- a rule is dropped and an image leaves its alt text behind.
"""

import io
from typing import TYPE_CHECKING

import docx
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from markdown_it import MarkdownIt
from markdown_it.tree import SyntaxTreeNode

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.paragraph import Paragraph

DOCX_MEDIA_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

_MAX_HEADING_LEVEL = 6
# Word's default template defines three levels of each list style; deeper nesting is clamped.
_MAX_LIST_LEVEL = 3
_CODE_FONT = 'Consolas'
# The default template has no Hyperlink character style, so links are coloured directly.
_LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)

_PARSER = MarkdownIt('commonmark', {'linkify': True}).enable(['table', 'strikethrough', 'linkify'])

# Inline nodes that switch a run property on for everything nested inside them.
_INLINE_FORMATS = {'strong': 'bold', 'em': 'italic', 's': 'strike'}


def _list_style(list_type: str, level: int) -> str:
    name = 'List Number' if list_type == 'ordered_list' else 'List Bullet'
    capped = min(level, _MAX_LIST_LEVEL - 1)
    return name if capped == 0 else f'{name} {capped + 1}'


def _add_run(paragraph: 'Paragraph', text: str, formats: frozenset[str]) -> None:
    # markdown-it brackets emphasis with empty text nodes; skipping them keeps the document
    # free of runs that carry formatting but no content.
    if not text:
        return

    run = paragraph.add_run(text)
    # Only ever switch formatting on: setting it to False would override the bold or italic a
    # paragraph style (a heading, say) already provides.
    if 'bold' in formats:
        run.bold = True
    if 'italic' in formats:
        run.italic = True
    if 'strike' in formats:
        run.font.strike = True
    if 'code' in formats:
        run.font.name = _CODE_FONT
    if 'link' in formats:
        run.font.color.rgb = _LINK_COLOR
        run.font.underline = True


def _write_link(paragraph: 'Paragraph', node: SyntaxTreeNode, formats: frozenset[str]) -> None:
    """Wrap the link's runs in a `w:hyperlink`, which python-docx has no public API for."""
    href = node.attrGet('href')
    already_written = len(paragraph.runs)
    _write_inline(paragraph, node, formats | {'link'})

    # `runs` only counts direct children, and earlier links have moved theirs into their own
    # hyperlink element, so whatever is past this point belongs to this link alone.
    new_runs = paragraph.runs[already_written:]
    if not isinstance(href, str) or not href or not new_runs:
        return

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), paragraph.part.relate_to(href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True))
    for run in new_runs:
        # Appending moves the run out of the paragraph and into the hyperlink.
        hyperlink.append(run._element)  # noqa: SLF001
    paragraph._p.append(hyperlink)  # noqa: SLF001


def _write_inline(paragraph: 'Paragraph', node: SyntaxTreeNode, formats: frozenset[str]) -> None:
    for child in node.children:
        if child.type in _INLINE_FORMATS:
            _write_inline(paragraph, child, formats | {_INLINE_FORMATS[child.type]})
        elif child.type == 'link':
            _write_link(paragraph, child, formats)
        elif child.type == 'code_inline':
            _add_run(paragraph, child.content, formats | {'code'})
        elif child.type in {'softbreak', 'hardbreak'}:
            _add_run(paragraph, ' ', formats)
        elif child.children:
            _write_inline(paragraph, child, formats)
        else:
            # Raw inline HTML and anything else unmapped is kept verbatim rather than dropped.
            _add_run(paragraph, child.content, formats)


def _write_paragraph(document: 'Document', node: SyntaxTreeNode, style: str | None) -> None:
    _write_inline(document.add_paragraph(style=style), node, frozenset())


def _write_list(document: 'Document', node: SyntaxTreeNode, level: int) -> None:
    style = _list_style(node.type, level)
    for item in node.children:
        for index, child in enumerate(item.children):
            if child.type in {'bullet_list', 'ordered_list'}:
                _write_list(document, child, level + 1)
            else:
                # Only the first block of an item is bulleted; the rest continue it.
                _write_block(document, child, style if index == 0 else 'List Continue')


def _write_code(document: 'Document', node: SyntaxTreeNode) -> None:
    for line in node.content.rstrip('\n').split('\n'):
        _add_run(document.add_paragraph(style='No Spacing'), line, frozenset({'code'}))


def _write_table(document: 'Document', node: SyntaxTreeNode) -> None:
    rows = [row for section in node.children for row in section.children]
    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=max(len(row.children) for row in rows))
    table.style = 'Table Grid'
    for row_index, row in enumerate(rows):
        for column_index, cell in enumerate(row.children):
            # Word has no header-cell concept, so GFM's `th` cells are emphasised instead.
            formats = frozenset({'bold'}) if cell.type == 'th' else frozenset()
            _write_inline(table.cell(row_index, column_index).paragraphs[0], cell, formats)


def _write_block(document: 'Document', node: SyntaxTreeNode, style: str | None = None) -> None:
    if node.type == 'heading':
        _write_paragraph(document, node, f'Heading {min(int(node.tag[1:]), _MAX_HEADING_LEVEL)}')
    elif node.type == 'paragraph':
        _write_paragraph(document, node, style)
    elif node.type in {'bullet_list', 'ordered_list'}:
        _write_list(document, node, level=0)
    elif node.type == 'blockquote':
        _write_blocks(document, node.children, 'Quote')
    elif node.type in {'fence', 'code_block'}:
        _write_code(document, node)
    elif node.type == 'table':
        _write_table(document, node)
    elif node.children:
        _write_blocks(document, node.children, style)
    elif node.content.strip():
        # Raw HTML is never interpreted; it is kept verbatim so no content disappears.
        document.add_paragraph(node.content.strip())


def _write_blocks(document: 'Document', nodes: list[SyntaxTreeNode], style: str | None = None) -> None:
    for node in nodes:
        _write_block(document, node, style)


def markdown_to_docx(markdown: str, *, title: str | None = None) -> bytes:
    """Render markdown as a .docx file and return its bytes."""
    document = docx.Document()
    if title:
        document.core_properties.title = title

    _write_blocks(document, SyntaxTreeNode(_PARSER.parse(markdown)).children)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

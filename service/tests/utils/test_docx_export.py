import io

import docx
import pytest
from docx.document import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ai_document_plugin_service.utils.docx_export import markdown_to_docx


def _render(markdown: str) -> Document:
    return docx.Document(io.BytesIO(markdown_to_docx(markdown)))


def _hyperlinks(paragraph: Paragraph) -> list:
    return paragraph._p.findall(qn('w:hyperlink'))  # noqa: SLF001


def _link_targets(paragraph: Paragraph) -> list[str]:
    """The external URL each hyperlink in the paragraph points at, in document order."""
    return [paragraph.part.rels[hyperlink.get(qn('r:id'))].target_ref for hyperlink in _hyperlinks(paragraph)]


def _hyperlink_runs(paragraph: Paragraph) -> list[Run]:
    return [Run(element, paragraph) for hyperlink in _hyperlinks(paragraph) for element in hyperlink.findall(qn('w:r'))]


def _hyperlink_text(paragraph: Paragraph) -> list[str]:
    return [
        ''.join(run.text for run in (Run(element, paragraph) for element in hyperlink.findall(qn('w:r'))))
        for hyperlink in _hyperlinks(paragraph)
    ]


def _style_name(paragraph: Paragraph) -> str | None:
    return paragraph.style.name if paragraph.style is not None else None


def _styles(document: Document) -> list[str | None]:
    return [_style_name(paragraph) for paragraph in document.paragraphs if paragraph.text]


def test_headings_use_word_heading_styles() -> None:
    document = _render('# Title\n\n## Section\n\n###### Deep')

    assert _styles(document) == ['Heading 1', 'Heading 2', 'Heading 6']
    assert [paragraph.text for paragraph in document.paragraphs if paragraph.text] == [
        'Title',
        'Section',
        'Deep',
    ]


def test_inline_formatting_maps_to_runs() -> None:
    document = _render('Plain **bold** and *italic* and ~~struck~~ and `code`.')

    runs = {run.text: run for run in document.paragraphs[0].runs}
    assert runs['bold'].bold is True
    assert runs['italic'].italic is True
    assert runs['struck'].font.strike is True
    assert runs['code'].font.name == 'Consolas'
    # Formatting is only ever switched on, never explicitly off, so styles keep inheriting.
    assert runs['Plain '].bold is None


def test_heading_runs_do_not_disable_inherited_bold() -> None:
    document = _render('# Title with **bold**')

    assert all(run.bold in (None, True) for run in document.paragraphs[0].runs)


def test_soft_break_becomes_a_space_in_one_paragraph() -> None:
    document = _render('first\nsecond')

    assert [paragraph.text for paragraph in document.paragraphs if paragraph.text] == ['first second']


def test_link_becomes_a_clickable_hyperlink() -> None:
    document = _render('See [the guide](https://example.org/guide) now.')

    paragraph = document.paragraphs[0]
    assert _link_targets(paragraph) == ['https://example.org/guide']
    assert _hyperlink_text(paragraph) == ['the guide']
    # The link is spliced in at the right position, not appended at the end of the paragraph.
    assert paragraph.text == 'See the guide now.'


def test_link_runs_are_styled_as_links() -> None:
    document = _render('[the guide](https://example.org/guide)')

    run = _hyperlink_runs(document.paragraphs[0])[0]
    assert run.font.underline is True
    assert run.font.color.rgb == RGBColor(0x05, 0x63, 0xC1)


def test_link_keeps_formatting_inside_its_text() -> None:
    document = _render('[**bold** link](https://example.org)')

    runs = _hyperlink_runs(document.paragraphs[0])
    assert [run.text for run in runs] == ['bold', ' link']
    assert [run.bold for run in runs] == [True, None]


def test_two_links_in_one_paragraph_keep_separate_targets() -> None:
    document = _render('[first](https://a.example) and [second](https://b.example)')

    paragraph = document.paragraphs[0]
    assert _link_targets(paragraph) == ['https://a.example', 'https://b.example']
    assert _hyperlink_text(paragraph) == ['first', 'second']


def test_link_inside_a_list_item_still_becomes_a_hyperlink() -> None:
    document = _render('- see [the guide](https://example.org/guide)')

    numbered = [p for p in document.paragraphs if _style_name(p) == 'List Bullet']
    assert _link_targets(numbered[0]) == ['https://example.org/guide']


def test_link_inside_a_table_cell_still_becomes_a_hyperlink() -> None:
    document = _render('| a |\n| --- |\n| [x](https://example.org) |')

    cell_paragraph = document.tables[0].cell(1, 0).paragraphs[0]
    assert _link_targets(cell_paragraph) == ['https://example.org']


def test_bare_url_becomes_a_hyperlink() -> None:
    # The pipeline writes URLs as bare text, not as [label](url), so linkify carries the
    # feature in practice. remark-gfm makes these clickable in the preview too.
    document = _render('See https://example.org/guide for details.')

    paragraph = document.paragraphs[0]
    assert _link_targets(paragraph) == ['https://example.org/guide']
    assert _hyperlink_text(paragraph) == ['https://example.org/guide']
    assert paragraph.text == 'See https://example.org/guide for details.'


def test_link_without_a_target_is_left_as_plain_text() -> None:
    document = _render('[just text]()')

    paragraph = document.paragraphs[0]
    assert _link_targets(paragraph) == []
    assert paragraph.text == 'just text'


def test_lists_use_built_in_list_styles() -> None:
    document = _render('- one\n- two\n\ntext\n\n1. first\n2. second')

    assert _styles(document) == [
        'List Bullet',
        'List Bullet',
        'Normal',
        'List Number',
        'List Number',
    ]


def test_nested_lists_step_through_the_style_levels() -> None:
    document = _render('- top\n    - nested\n        - deeper')

    assert _styles(document) == ['List Bullet', 'List Bullet 2', 'List Bullet 3']


def test_list_nesting_deeper_than_three_levels_is_clamped() -> None:
    markdown = '\n'.join(f'{" " * (4 * depth)}- level {depth}' for depth in range(6))
    document = _render(markdown)

    assert _styles(document)[-3:] == ['List Bullet 3', 'List Bullet 3', 'List Bullet 3']


def test_loose_list_item_continuation_is_not_bulleted_again() -> None:
    document = _render('- first para\n\n  second para\n\n- next item')

    assert _styles(document) == ['List Bullet', 'List Continue', 'List Bullet']


def test_consecutive_ordered_lists_share_one_counter() -> None:
    """Documents a known limitation of mapping onto Word's built-in `List Number` style.

    Word gives that style a single numbering definition, so a second list continues 3., 4.
    rather than restarting at 1. Restarting would mean writing numbering XML by hand, which
    this module deliberately avoids. The test exists so the behaviour is a recorded choice
    rather than a surprise.
    """
    document = _render('1. one\n\ntext\n\n1. alpha')

    assert _styles(document) == ['List Number', 'Normal', 'List Number']


def test_table_renders_with_bold_header() -> None:
    document = _render('| Name | Size |\n| --- | ---: |\n| disk | 10 |\n| tape | 20 |')

    table = document.tables[0]
    assert (len(table.rows), len(table.columns)) == (3, 2)
    assert table.cell(0, 0).text == 'Name'
    assert table.cell(2, 1).text == '20'
    assert all(run.bold for run in table.cell(0, 0).paragraphs[0].runs)
    assert not any(run.bold for run in table.cell(1, 0).paragraphs[0].runs)


def test_table_cell_keeps_inline_formatting() -> None:
    document = _render('| a |\n| --- |\n| **strong** |')

    assert document.tables[0].cell(1, 0).paragraphs[0].runs[0].bold is True


def test_fenced_code_keeps_one_paragraph_per_line() -> None:
    document = _render('```python\nfirst = 1\nsecond = 2\n```')

    code = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.runs and paragraph.runs[0].font.name == 'Consolas'
    ]
    assert code == ['first = 1', 'second = 2']


def test_raw_html_is_kept_as_literal_text_not_interpreted() -> None:
    document = _render('<script>alert(1)</script>')

    assert 'alert(1)' in '\n'.join(paragraph.text for paragraph in document.paragraphs)


def test_unsupported_constructs_degrade_without_breaking_the_document() -> None:
    """Images and thematic breaks are deliberately not handled; the pipeline never emits them.

    They must still not raise or corrupt the file if one ever reaches the editor: a rule is
    dropped and an image leaves its alt text behind.
    """
    document = _render('above\n\n---\n\n![a diagram](https://example.org/x.png)\n\nbelow')

    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert texts == ['above', 'a diagram', 'below']


def test_title_is_stored_as_document_metadata() -> None:
    document = docx.Document(io.BytesIO(markdown_to_docx('# Body', title='My Plan')))

    assert document.core_properties.title == 'My Plan'


@pytest.mark.parametrize('markdown', ['', '   \n\n  '])
def test_empty_markdown_still_produces_a_readable_file(markdown: str) -> None:
    document = _render(markdown)

    assert all(not paragraph.text for paragraph in document.paragraphs)

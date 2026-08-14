"""Convert generated markdown into a Word (.docx) document.

The markdown produced by the pipeline is CommonMark plus the GitHub extensions the plugin
renders in the browser (tables, strikethrough), so the parser here is configured to match
`remark-gfm` on the frontend. Every construct is mapped onto Word's built-in styles rather
than hand-tuned formatting, so headings feed Word's navigation pane and generated tables of
contents, and so an institution can restyle the output by editing the styles.
"""

import io
from dataclasses import dataclass, replace
from typing import TYPE_CHECKING

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.tree import SyntaxTreeNode

if TYPE_CHECKING:
    from docx.document import Document
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

DOCX_MEDIA_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

_MAX_HEADING_LEVEL = 6
_MAX_LIST_LEVEL = 5
_LEVELS_PER_ABSTRACT_NUM = _MAX_LIST_LEVEL + 1

# Word measures indentation in twips (1/20 pt); 720 twips is the usual half-inch list step.
_INDENT_STEP_TWIPS = 720
_HANGING_INDENT_TWIPS = 360

# IDs for the numbering definitions this module adds. Kept high so they cannot collide with
# whatever the python-docx default template already defines.
_NUMBERING_ID_FLOOR = 9000

_CODE_FONT = 'Consolas'
_CODE_FONT_SIZE_PT = 9.5
_CODE_INDENT_PT = 18
_LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)

_TABLE_STYLE = 'Table Grid'
_QUOTE_STYLE = 'Quote'
_LIST_PARAGRAPH_STYLE = 'List Paragraph'
_CODE_PARAGRAPH_STYLE = 'No Spacing'

# Bullet glyphs Word itself uses for the first three list levels, then repeating. The first and
# third are private-use code points (U+F0B7 in Symbol, U+F0A7 in Wingdings), which is how Word
# encodes them; they only render correctly alongside the w:rFonts below.
_BULLET_GLYPHS = (('', 'Symbol'), ('o', 'Courier New'), ('', 'Wingdings'))
_ORDERED_FORMATS = ('decimal', 'lowerLetter', 'lowerRoman')

_CELL_ALIGNMENTS = {
    'text-align:left': WD_ALIGN_PARAGRAPH.LEFT,
    'text-align:center': WD_ALIGN_PARAGRAPH.CENTER,
    'text-align:right': WD_ALIGN_PARAGRAPH.RIGHT,
}


@dataclass(frozen=True)
class _RunSpec:
    """A stretch of text together with the inline formatting that applies to it."""

    text: str = ''
    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False
    link: str | None = None
    is_break: bool = False


def _build_parser() -> MarkdownIt:
    """CommonMark plus the GFM extensions the frontend renders via remark-gfm."""
    return MarkdownIt('commonmark').enable(['table', 'strikethrough'])


def _has_style(document: 'Document', name: str) -> bool:
    try:
        document.styles[name]
    except KeyError:
        return False
    return True


def _nested_format(node: SyntaxTreeNode, base: _RunSpec) -> _RunSpec | None:
    """The formatting a wrapping node adds to its children, or None if it does not wrap."""
    if node.type == 'strong':
        return replace(base, bold=True)
    if node.type == 'em':
        return replace(base, italic=True)
    if node.type == 's':
        return replace(base, strike=True)
    if node.type == 'link':
        href = node.attrGet('href')
        return replace(base, link=href if isinstance(href, str) else None)
    return None


def _leaf_spec(node: SyntaxTreeNode, base: _RunSpec) -> _RunSpec | None:
    """The run a self-contained inline node produces, or None if it is not one."""
    if node.type == 'text':
        return replace(base, text=node.content)
    if node.type == 'code_inline':
        return replace(base, text=node.content, code=True)
    if node.type == 'softbreak':
        return replace(base, text=' ')
    if node.type == 'hardbreak':
        return replace(base, text='', is_break=True)
    if node.type == 'image':
        # Images are not embedded: the markdown only carries a URL, and fetching remote media
        # server-side would turn an export into an outbound request.
        alt = node.content or node.attrGet('alt') or 'image'
        return replace(base, text=f'[{alt}]')
    return None


def _flatten_inline(node: SyntaxTreeNode, base: _RunSpec, out: list[_RunSpec]) -> None:
    """Flatten an inline subtree into a linear list of formatted runs."""
    for child in node.children:
        nested = _nested_format(child, base)
        if nested is not None:
            _flatten_inline(child, nested, out)
            continue

        leaf = _leaf_spec(child, base)
        if leaf is not None:
            out.append(leaf)
        elif child.children:
            _flatten_inline(child, base, out)
        elif child.content:
            # Anything unmapped (raw HTML, footnote markers) is kept as literal text so no
            # content is silently dropped.
            out.append(replace(base, text=child.content))


def _inline_specs(node: SyntaxTreeNode) -> list[_RunSpec]:
    """Collect the run specs for a block node whose first child is an inline node."""
    specs: list[_RunSpec] = []
    for child in node.children:
        if child.type == 'inline':
            _flatten_inline(child, _RunSpec(), specs)
    return specs


def _add_run(paragraph: 'Paragraph', spec: _RunSpec) -> 'Run':
    run = paragraph.add_run()
    if spec.is_break:
        run.add_break()
        return run

    run.text = spec.text
    # Only ever turn formatting on: setting these to False would override the bold/italic a
    # paragraph style (a heading, say) already provides.
    if spec.bold:
        run.bold = True
    if spec.italic:
        run.italic = True
    if spec.strike:
        run.font.strike = True
    if spec.code:
        run.font.name = _CODE_FONT
        run.font.size = Pt(_CODE_FONT_SIZE_PT)
    return run


def _add_hyperlink(paragraph: 'Paragraph', url: str, specs: list[_RunSpec]) -> None:
    """Wrap runs in a `w:hyperlink`; python-docx has no public API for external links."""
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), relationship_id)

    for spec in specs:
        run = _add_run(paragraph, spec)
        run.font.color.rgb = _LINK_COLOR
        run.font.underline = True
        # Appending moves the run out of the paragraph and into the hyperlink element.
        hyperlink.append(run._element)  # noqa: SLF001

    paragraph._p.append(hyperlink)  # noqa: SLF001


def _write_specs(paragraph: 'Paragraph', specs: list[_RunSpec]) -> None:
    index = 0
    while index < len(specs):
        link = specs[index].link
        if link is None:
            _add_run(paragraph, specs[index])
            index += 1
            continue

        group: list[_RunSpec] = []
        while index < len(specs) and specs[index].link == link:
            group.append(specs[index])
            index += 1
        _add_hyperlink(paragraph, link, group)


def _bullet_level_xml(level: int) -> str:
    glyph, font = _BULLET_GLYPHS[level % len(_BULLET_GLYPHS)]
    indent = _INDENT_STEP_TWIPS * (level + 1)
    return (
        f'<w:lvl w:ilvl="{level}">'
        f'<w:start w:val="1"/>'
        f'<w:numFmt w:val="bullet"/>'
        f'<w:lvlText w:val="{glyph}"/>'
        f'<w:lvlJc w:val="left"/>'
        f'<w:pPr><w:ind w:left="{indent}" w:hanging="{_HANGING_INDENT_TWIPS}"/></w:pPr>'
        f'<w:rPr><w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:hint="default"/></w:rPr>'
        f'</w:lvl>'
    )


def _ordered_level_xml(level: int, start: int) -> str:
    number_format = _ORDERED_FORMATS[level % len(_ORDERED_FORMATS)]
    indent = _INDENT_STEP_TWIPS * (level + 1)
    return (
        f'<w:lvl w:ilvl="{level}">'
        f'<w:start w:val="{start}"/>'
        f'<w:numFmt w:val="{number_format}"/>'
        f'<w:lvlText w:val="%{level + 1}."/>'
        f'<w:lvlJc w:val="left"/>'
        f'<w:pPr><w:ind w:left="{indent}" w:hanging="{_HANGING_INDENT_TWIPS}"/></w:pPr>'
        f'</w:lvl>'
    )


def _write_cell(cell: '_Cell', node: SyntaxTreeNode) -> None:
    paragraph = cell.paragraphs[0]

    style = node.attrGet('style')
    if isinstance(style, str):
        alignment = _CELL_ALIGNMENTS.get(style.replace(' ', ''))
        if alignment is not None:
            paragraph.alignment = alignment

    # GFM header cells are `th`; Word has no header-cell concept, so emphasise the text.
    specs = _inline_specs(node)
    if node.type == 'th':
        specs = [replace(spec, bold=True) for spec in specs]
    _write_specs(paragraph, specs)


class _NumberingRegistry:
    """Allocates a fresh numbering definition per list.

    Reusing Word's `List Number` style would make every ordered list in the document share a
    single counter, so a second list would continue 4., 5., 6. instead of restarting. Giving
    each list its own definition also lets nested levels carry their own indentation.
    """

    def __init__(self, document: 'Document') -> None:
        self._numbering = document.part.numbering_part.element
        self._next_id = self._first_free_id()

    def _first_free_id(self) -> int:
        used = [_NUMBERING_ID_FLOOR]
        for tag, attribute in (('w:abstractNum', 'w:abstractNumId'), ('w:num', 'w:numId')):
            for element in self._numbering.findall(qn(tag)):
                value = element.get(qn(attribute))
                if value is not None and value.isdigit():
                    used.append(int(value))
        return max(used) + 1

    def create(self, *, ordered: bool, start: int) -> int:
        """Define a new list and return the `numId` its paragraphs should reference."""
        abstract_num_id = self._next_id
        num_id = self._next_id + 1
        self._next_id += 2

        levels = ''.join(
            _ordered_level_xml(level, start if level == 0 else 1) if ordered else _bullet_level_xml(level)
            for level in range(_LEVELS_PER_ABSTRACT_NUM)
        )
        abstract_num = parse_xml(
            f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abstract_num_id}">'
            f'<w:multiLevelType w:val="hybridMultilevel"/>'
            f'{levels}'
            f'</w:abstractNum>'
        )
        num = parse_xml(
            f'<w:num {nsdecls("w")} w:numId="{num_id}"><w:abstractNumId w:val="{abstract_num_id}"/></w:num>'
        )

        # The schema requires every w:abstractNum to precede every w:num.
        existing_nums = self._numbering.findall(qn('w:num'))
        if existing_nums:
            existing_nums[0].addprevious(abstract_num)
        else:
            self._numbering.append(abstract_num)
        self._numbering.append(num)

        return num_id


def _apply_numbering(paragraph: 'Paragraph', num_id: int, level: int) -> None:
    properties = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    numbering = properties.get_or_add_numPr()
    numbering.get_or_add_ilvl().val = level
    numbering.get_or_add_numId().val = num_id


class _DocxWriter:
    def __init__(self, document: 'Document') -> None:
        self._document = document
        self._numbering = _NumberingRegistry(document)
        self._has_quote_style = _has_style(document, _QUOTE_STYLE)
        self._has_list_style = _has_style(document, _LIST_PARAGRAPH_STYLE)
        self._has_code_style = _has_style(document, _CODE_PARAGRAPH_STYLE)
        self._has_table_style = _has_style(document, _TABLE_STYLE)

    def write(self, nodes: list[SyntaxTreeNode]) -> None:
        for node in nodes:
            self._write_block(node, indent_level=0)

    def _new_paragraph(self, *, style: str | None = None, indent_level: int = 0) -> 'Paragraph':
        paragraph = self._document.add_paragraph(style=style)
        if indent_level > 0:
            paragraph.paragraph_format.left_indent = Pt(_CODE_INDENT_PT * indent_level)
        return paragraph

    def _write_block(self, node: SyntaxTreeNode, *, indent_level: int) -> None:
        if node.type == 'heading':
            self._write_heading(node, indent_level=indent_level)
        elif node.type == 'paragraph':
            paragraph = self._new_paragraph(indent_level=indent_level)
            _write_specs(paragraph, _inline_specs(node))
        elif node.type in {'bullet_list', 'ordered_list'}:
            self._write_list(node, level=0, indent_level=indent_level)
        elif node.type == 'blockquote':
            self._write_blockquote(node, indent_level=indent_level)
        else:
            self._write_standalone_block(node, indent_level=indent_level)

    def _write_standalone_block(self, node: SyntaxTreeNode, *, indent_level: int) -> None:
        if node.type in {'fence', 'code_block'}:
            self._write_code(node, indent_level=indent_level)
        elif node.type == 'table':
            self._write_table(node)
        elif node.type == 'hr':
            self._write_rule()
        elif node.type == 'html_block':
            # Raw HTML is never interpreted: it is written verbatim so nothing disappears.
            paragraph = self._new_paragraph(indent_level=indent_level)
            paragraph.add_run(node.content.strip())
        elif node.children:
            for child in node.children:
                self._write_block(child, indent_level=indent_level)

    def _write_heading(self, node: SyntaxTreeNode, *, indent_level: int) -> None:
        level = min(int(node.tag[1:]), _MAX_HEADING_LEVEL)
        paragraph = self._new_paragraph(style=f'Heading {level}', indent_level=indent_level)
        _write_specs(paragraph, _inline_specs(node))

    def _write_blockquote(self, node: SyntaxTreeNode, *, indent_level: int) -> None:
        for child in node.children:
            if child.type == 'paragraph' and self._has_quote_style:
                paragraph = self._new_paragraph(style=_QUOTE_STYLE, indent_level=indent_level)
                _write_specs(paragraph, _inline_specs(child))
            else:
                self._write_block(child, indent_level=indent_level + 1)

    def _write_code(self, node: SyntaxTreeNode, *, indent_level: int) -> None:
        style = _CODE_PARAGRAPH_STYLE if self._has_code_style else None
        for line in node.content.rstrip('\n').split('\n'):
            paragraph = self._new_paragraph(style=style, indent_level=max(indent_level, 1))
            run = paragraph.add_run(line)
            run.font.name = _CODE_FONT
            run.font.size = Pt(_CODE_FONT_SIZE_PT)

    def _write_rule(self) -> None:
        paragraph = self._document.add_paragraph()
        properties = paragraph._p.get_or_add_pPr()  # noqa: SLF001
        # Safe to append: the paragraph carries no other properties, so w:pBdr cannot land
        # out of schema order.
        properties.append(
            parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
        )

    def _write_list(self, node: SyntaxTreeNode, *, level: int, indent_level: int) -> None:
        ordered = node.type == 'ordered_list'
        start_attr = node.attrGet('start') if ordered else None
        start = int(start_attr) if start_attr is not None and str(start_attr).isdigit() else 1
        num_id = self._numbering.create(ordered=ordered, start=start)
        self._write_list_items(node, num_id=num_id, level=level, indent_level=indent_level)

    def _write_list_items(self, node: SyntaxTreeNode, *, num_id: int, level: int, indent_level: int) -> None:
        capped_level = min(level, _MAX_LIST_LEVEL)
        style = _LIST_PARAGRAPH_STYLE if self._has_list_style else None

        for item in node.children:
            if item.type != 'list_item':
                continue

            numbered = False
            for child in item.children:
                if child.type in {'bullet_list', 'ordered_list'}:
                    # A nested list of the same kind stays on the parent's counter, so that
                    # a.  b.  c. continues correctly across the nesting.
                    if child.type == node.type:
                        self._write_list_items(child, num_id=num_id, level=level + 1, indent_level=indent_level)
                    else:
                        self._write_list(child, level=level + 1, indent_level=indent_level)
                elif child.type == 'paragraph' and not numbered:
                    paragraph = self._new_paragraph(style=style, indent_level=indent_level)
                    _apply_numbering(paragraph, num_id, capped_level)
                    _write_specs(paragraph, _inline_specs(child))
                    numbered = True
                else:
                    # Continuation blocks in a loose item: indented, but not re-numbered.
                    self._write_block(child, indent_level=indent_level + capped_level + 1)

    def _write_table(self, node: SyntaxTreeNode) -> None:
        rows = [row for section in node.children for row in section.children if row.type == 'tr']
        if not rows:
            return

        column_count = max(len(row.children) for row in rows)
        table = self._document.add_table(rows=len(rows), cols=column_count)
        if self._has_table_style:
            table.style = _TABLE_STYLE

        for row_index, row in enumerate(rows):
            for column_index, cell_node in enumerate(row.children):
                if column_index >= column_count:
                    continue
                _write_cell(table.cell(row_index, column_index), cell_node)


def markdown_to_docx(markdown: str, *, title: str | None = None) -> bytes:
    """Render markdown as a .docx file and return its bytes."""
    tree = SyntaxTreeNode(_build_parser().parse(markdown))

    document = docx.Document()
    if title:
        document.core_properties.title = title

    _DocxWriter(document).write(list(tree.children))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
